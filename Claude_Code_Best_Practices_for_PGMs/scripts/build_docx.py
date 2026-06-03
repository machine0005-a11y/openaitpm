#!/usr/bin/env python3
"""Build the Claude Code Best Practices for PgMs visual guide DOCX."""

from __future__ import annotations

import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
except ModuleNotFoundError as exc:
    raise SystemExit(
        "Missing dependency: python-docx. Install it with "
        "`python3 -m pip install python-docx`, then rerun this script."
    ) from exc


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = PROJECT_ROOT / "source" / "guide_content.md"
OUTPUT_PATH = PROJECT_ROOT / "output" / "Claude_Code_Best_Practices_for_PGMs_Visual_Guide.docx"

COLORS = {
    "title": "17365D",
    "heading": "1F4E79",
    "subtitle": "4F81BD",
    "caption": "595959",
    "footer": "7F7F7F",
    "white": "FFFFFF",
    "table_header": "1F4E79",
    "source_header": "5B9BD5",
    "decision_header": "404040",
    "table_body": "F7F9FB",
    "border": "D9E2EC",
    "gold": "FFF2CC",
    "green": "E2F0D9",
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def parse_frontmatter(lines: list[str]) -> tuple[dict[str, str], int]:
    if not lines or lines[0].strip() != "---":
        return {}, 0
    meta: dict[str, str] = {}
    idx = 1
    while idx < len(lines):
        line = lines[idx].strip()
        if line == "---":
            return meta, idx + 1
        if ":" in line:
            key, value = line.split(":", 1)
            meta[key.strip()] = value.strip()
        idx += 1
    return meta, idx


def ensure_paragraph_style(doc: Document, name: str):
    styles = doc.styles
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(26)
    title.font.color.rgb = rgb(COLORS["title"])
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Aptos"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = rgb(COLORS["heading"])
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    subtitle = ensure_paragraph_style(doc, "Guide Subtitle")
    subtitle.base_style = normal
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(12)
    subtitle.font.italic = True
    subtitle.font.color.rgb = rgb(COLORS["subtitle"])
    subtitle.paragraph_format.space_after = Pt(10)

    caption = ensure_paragraph_style(doc, "Guide Caption")
    caption.base_style = normal
    caption.font.name = "Aptos"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(COLORS["caption"])
    caption.paragraph_format.space_after = Pt(8)

    prompt = ensure_paragraph_style(doc, "Prompt Quote")
    prompt.base_style = normal
    prompt.font.name = "Aptos"
    prompt.font.size = Pt(10)
    prompt.font.bold = True
    prompt.font.italic = True
    prompt.font.color.rgb = rgb(COLORS["subtitle"])
    prompt.paragraph_format.left_indent = Inches(0.35)
    prompt.paragraph_format.space_before = Pt(2)
    prompt.paragraph_format.space_after = Pt(10)


def configure_page(doc: Document, footer_text: str) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = ""
    run = paragraph.add_run(footer_text)
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["footer"])


def apply_run_format(run, *, bold: bool | None = None, italic: bool | None = None,
                     size=None, color: str | None = None) -> None:
    run.font.name = "Aptos"
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = size
    if color:
        run.font.color.rgb = rgb(color)


def add_markdown_runs(paragraph, text: str, *, bold: bool = False,
                      italic: bool = False, size=None, color: str | None = None) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            apply_run_format(run, bold=bold, italic=italic, size=size, color=color)
        run = paragraph.add_run(match.group(1))
        apply_run_format(run, bold=True, italic=italic, size=size, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        apply_run_format(run, bold=bold, italic=italic, size=size, color=color)


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None,
                            *, alignment=None) -> None:
    paragraph = doc.add_paragraph(style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    add_markdown_runs(paragraph, text)


def add_image_or_placeholder(doc: Document, source_path: Path, alt_text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    if source_path.exists():
        run.add_picture(str(source_path), width=Inches(6.9))
        return
    add_markdown_runs(paragraph, f"[Missing figure asset: {alt_text}]",
                      bold=True, color=COLORS["caption"])


def get_or_add(parent, tag: str):
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    return element


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = get_or_add(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    for tag, value in (("w:top", top), ("w:start", start), ("w:bottom", bottom), ("w:end", end)):
        node = get_or_add(tc_mar, tag)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = get_or_add(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = get_or_add(tbl_pr, "w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = COLORS["border"]) -> None:
    tbl_pr = table._tbl.tblPr
    borders = get_or_add(tbl_pr, "w:tblBorders")
    for edge in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        node = get_or_add(borders, edge)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def table_profile(headers: list[str]) -> tuple[list[float], str]:
    header_key = tuple(headers)
    if header_key == ("Practice", "Why It Matters", "Prompt / Action"):
        return [2.0, 1.9, 3.0], COLORS["table_header"]
    if header_key == ("Source", "Use It For", "Guardrail"):
        return [1.35, 2.75, 2.8], COLORS["source_header"]
    if header_key == ("Date", "Decision", "Why", "Tradeoff / Risk", "Owner"):
        return [1.0, 1.7, 1.4, 2.0, 0.8], COLORS["decision_header"]
    return [6.9 / len(headers)] * len(headers), COLORS["table_header"]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    widths, header_fill = table_profile(headers)
    table = doc.add_table(rows=len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                add_markdown_runs(paragraph, value, bold=True, size=Pt(9), color=COLORS["white"])
            else:
                set_cell_shading(cell, COLORS["table_body"])
                add_markdown_runs(paragraph, value, size=Pt(9))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def strip_bold_markup(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text).strip()


def add_callout(doc: Document, lines: list[str], kind: str) -> None:
    fill = COLORS.get(kind, COLORS["gold"])
    cleaned = [line.strip() for line in lines if line.strip()]
    if not cleaned:
        return

    title = strip_bold_markup(cleaned[0])
    body = "\n".join(cleaned[1:])
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.9])
    set_table_borders(table, "FFFFFF")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    title_paragraph = cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(2)
    add_markdown_runs(title_paragraph, title, bold=True)

    if body:
        body_paragraph = cell.add_paragraph()
        body_paragraph.paragraph_format.space_after = Pt(0)
        add_markdown_runs(body_paragraph, body)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def is_table_divider(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not is_table_divider(cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def render_markdown(doc: Document, source_text: str, source_dir: Path) -> None:
    lines = source_text.splitlines()
    metadata, idx = parse_frontmatter(lines)
    footer_text = metadata.get("footer", "Claude Code Best Practices for PgMs | Draft v1")
    configure_page(doc, footer_text)

    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.strip()
        idx += 1

        if not line:
            continue

        if line.startswith("::: callout"):
            parts = line.split()
            kind = parts[2] if len(parts) >= 3 else "gold"
            block: list[str] = []
            while idx < len(lines) and lines[idx].strip() != ":::":
                block.append(lines[idx])
                idx += 1
            idx += 1
            add_callout(doc, block, kind)
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", line)
        if image_match:
            alt_text, target = image_match.groups()
            image_path = (source_dir / target).resolve()
            add_image_or_placeholder(doc, image_path, alt_text)
            continue

        if line.startswith("|"):
            rows, idx = parse_table(lines, idx - 1)
            add_markdown_table(doc, rows)
            continue

        if line.startswith("# "):
            add_formatted_paragraph(doc, line[2:].strip(), style="Title")
            continue

        if line.startswith("## "):
            paragraph = doc.add_heading(line[3:].strip(), level=1)
            paragraph.paragraph_format.keep_with_next = True
            continue

        if line.startswith("_") and line.endswith("_"):
            add_formatted_paragraph(doc, line.strip("_"), style="Guide Subtitle")
            continue

        if line.startswith("*Figure") and line.endswith("*"):
            add_formatted_paragraph(doc, line.strip("*"), style="Guide Caption")
            continue

        if line.startswith("> "):
            add_formatted_paragraph(doc, line[2:].strip(), style="Prompt Quote")
            continue

        numbered = re.match(r"\d+\.\s+(.+)", line)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_markdown_runs(paragraph, numbered.group(1).strip())
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(paragraph, line[2:].strip())
            continue

        if line.startswith("**") and line.endswith("**"):
            add_formatted_paragraph(doc, line, style="Normal")
            continue

        add_formatted_paragraph(doc, line, style="Normal")


def build_docx() -> Path:
    source_text = SOURCE_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    render_markdown(doc, source_text, SOURCE_PATH.parent)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    output_path = build_docx()
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
